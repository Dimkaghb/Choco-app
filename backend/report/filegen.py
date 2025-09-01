
import docx
from docx import Document


Doc = Document()


Doc.add_paragraph('Hello, World!')
Doc.add_heading('My Heading', level=1)
Doc.add_heading('My Subheading', level=2)



Doc.save('report.docx')
